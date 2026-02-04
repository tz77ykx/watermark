import docx
import fitz

def create_docx():
    doc = docx.Document()
    doc.add_heading('Sample Exam Paper', 0)
    for i in range(10):
        doc.add_paragraph(f"This is paragraph {i}. " * 30)
    doc.save("sample.docx")
    print("Created sample.docx")

def create_pdf():
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), "Sample PDF Exam Paper\n" + ("Line of text.\n" * 20), fontsize=12)
    doc.save("sample.pdf")
    print("Created sample.pdf")

if __name__ == "__main__":
    create_docx()
    create_pdf()
